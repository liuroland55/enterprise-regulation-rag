"""文档文本抽取模块（Document Text Extraction）。

本模块在**不修改 RAG 核心 / ingestion loader**（`src/ingestion/loader.py` 等）的前提下，
为侧车服务提供对**富文本文档（PDF / DOCX）**的文本抽取能力，作为 loader 仅支持
UTF-8 纯文本（.txt/.md/.rst/.log）的**附加式增强**：

- 纯文本扩展名沿用 loader 的口径，直接以 utf-8 读取；
- `.pdf` 经 `pypdf` 逐页抽取并拼接；
- `.docx` 经 `python-docx` 抽取全部段落与表格单元格文本并拼接；
- `.doc`（旧版二进制 Word）**有意不支持**，抛出 `UnsupportedDocError` 并提示转换；
- 产出的 `langchain_core` `Document` 与 loader 写入的元数据口径保持一致
  （source / filepath / filetype / size），从而可直接交给 `loader.split_documents`。

设计约束：
- **零核心改动**：本模块不导入 / 不修改 loader 内部实现；仅在调用方处复用
  `loader.split_documents` / `loader.get_vector_store`（由 router / indexing 负责）。
- **惰性 / 防御式导入**：`pypdf` / `docx` 在各分支内部延迟导入，避免模块加载期引入
  这些可选依赖；缺失或解析失败时以清晰错误或空串兜底，绝不让调用方崩溃。

约定：代码 / API 名称使用英文；注释使用中文。
"""

from __future__ import annotations

import logging
import os
from typing import List, Optional

from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# 纯文本扩展名集合：与 loader.SUPPORTED_EXTENSIONS 同口径（以 utf-8 读取）。
TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".log"}

# 富文本扩展名集合：本模块新增支持的二进制 / 结构化文档类型。
RICH_EXTENSIONS = {".pdf", ".docx"}

# 侧车服务对外声明的“受支持扩展名”全集（注意：.doc 有意不在其中）。
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | RICH_EXTENSIONS


class UnsupportedDocError(Exception):
    """明确不受支持的文档类型（如旧版 .doc）所抛出的异常。

    与“解析失败 / 内容为空”区分开来：本异常代表“该类型本就不支持”，
    调用方（如上传接口）可据此返回 HTTP 415 并提示用户转换格式。
    """


def _extract_pdf(filepath: str) -> str:
    """抽取 PDF 全部页面文本并以换行拼接。

    使用 `pypdf.PdfReader`，逐页调用 `page.extract_text()`（可能返回 None，需防御）。
    解析失败时记录日志并抛出异常，由上层统一兜底。
    """
    from pypdf import PdfReader  # 惰性导入：仅在处理 PDF 时引入

    reader = PdfReader(filepath)
    parts: List[str] = []
    for page in reader.pages:
        # extract_text() 在无文本层（如纯扫描件）时可能返回 None
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts)


def _extract_docx(filepath: str) -> str:
    """抽取 DOCX 的全部段落与表格单元格文本，并以换行拼接。

    使用 `python-docx`（导入名为 `docx`）。先收集所有段落文本，再遍历所有表格的
    每个单元格文本，保证表格中的内容也能进入知识库。
    """
    import docx  # 惰性导入：仅在处理 DOCX 时引入（python-docx 包名为 docx）

    document = docx.Document(filepath)
    parts: List[str] = []

    # 1) 段落文本
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)

    # 2) 表格单元格文本（按行、按单元格顺序）
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def extract_text(filepath: str) -> str:
    """按扩展名分派抽取文档纯文本，返回拼接后的字符串。

    分派规则（扩展名一律小写）：
    - 纯文本（.txt/.md/.rst/.log）：以 utf-8 打开（errors="ignore"）读取；
    - .pdf：经 pypdf 逐页抽取；
    - .docx：经 python-docx 抽取段落 + 表格；
    - .doc：抛出 `UnsupportedDocError`（提示转换为 .docx 或 PDF）；
    - 其它：抛出 `UnsupportedDocError`。

    健壮性：对富文本解析包裹 try/except，解析失败时抛出清晰错误（由上层 extract_single
    兜底为 None）；空内容返回空串而非崩溃。
    """
    ext = os.path.splitext(filepath)[1].lower()

    # 纯文本：直接 utf-8 读取（errors="ignore" 容忍个别非法字节）
    if ext in TEXT_EXTENSIONS:
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as exc:  # noqa: BLE001 - 读取失败时抛出清晰错误，交由上层兜底
            raise RuntimeError(f"Failed to read text file '{filepath}': {exc}") from exc

    # PDF
    if ext == ".pdf":
        try:
            return _extract_pdf(filepath)
        except Exception as exc:  # noqa: BLE001 - pypdf 缺失 / 解析失败：抛出清晰错误
            raise RuntimeError(f"Failed to extract PDF '{filepath}': {exc}") from exc

    # DOCX
    if ext == ".docx":
        try:
            return _extract_docx(filepath)
        except Exception as exc:  # noqa: BLE001 - python-docx 缺失 / 解析失败：抛出清晰错误
            raise RuntimeError(f"Failed to extract DOCX '{filepath}': {exc}") from exc

    # 旧版 .doc：有意不支持，提示转换
    if ext == ".doc":
        raise UnsupportedDocError(
            "Legacy .doc is not supported; please convert to .docx or PDF."
        )

    # 其它一切类型：不支持
    raise UnsupportedDocError(f"Unsupported document type: {ext or '(none)'}")


def extract_single(filepath: str) -> Optional[Document]:
    """将单个文件抽取为一个 `langchain_core` `Document`。

    - `page_content` 取 `extract_text(filepath)`；
    - `metadata` 与 loader 写入口径一致：``source``（basename）/ ``filepath`` /
      ``filetype``（小写扩展名）/ ``size``（抽取文本的字符数）。

    返回 None 的情况（记录 warning，绝不抛出）：
    - 抽取文本为空或仅空白；
    - 抽取过程抛出任何异常（含 `UnsupportedDocError`）。
    """
    basename = os.path.basename(filepath)
    ext = os.path.splitext(basename)[1].lower()

    try:
        text = extract_text(filepath)
    except UnsupportedDocError as exc:
        # 不支持的类型：跳过并记录（上层若需 415 可自行先行校验扩展名）
        logger.warning("extract_single: unsupported document '%s': %s", basename, exc)
        return None
    except Exception:  # noqa: BLE001 - 抽取失败：跳过该文件，不影响其它文件
        logger.warning("extract_single: failed to extract '%s'; skipped.", basename, exc_info=True)
        return None

    # 空内容（含仅空白）：视为无可索引内容，跳过
    if not text or not text.strip():
        logger.warning("extract_single: empty content extracted from '%s'; skipped.", basename)
        return None

    return Document(
        page_content=text,
        metadata={
            "source": basename,
            "filepath": filepath,
            "filetype": ext,
            "size": len(text),
        },
    )


def load_documents_any(directory: str) -> List[Document]:
    """枚举目录中受支持的文件并抽取为 `Document` 列表（富文本版的 loader 类比）。

    - 仅纳入扩展名 ∈ `SUPPORTED_EXTENSIONS` 的常规文件；
    - 对每个文件调用 `extract_single`，收集非 None 的结果；
    - 单个文件失败 / 为空 / 不支持时跳过（由 `extract_single` 记录 warning）。

    这是 `loader.load_documents_from_directory` 的富格式类比，供 reindex / sync 复用。
    """
    documents: List[Document] = []

    if not os.path.isdir(directory):
        logger.warning("load_documents_any: directory '%s' is not a valid directory.", directory)
        return documents

    for name in sorted(os.listdir(directory)):
        path = os.path.join(directory, name)
        if not os.path.isfile(path):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        doc = extract_single(path)
        if doc is not None:
            documents.append(doc)

    return documents


__all__ = [
    "TEXT_EXTENSIONS",
    "RICH_EXTENSIONS",
    "SUPPORTED_EXTENSIONS",
    "UnsupportedDocError",
    "extract_text",
    "extract_single",
    "load_documents_any",
]
